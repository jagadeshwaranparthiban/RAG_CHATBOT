import io
from pathlib import Path
from typing import List
from uuid import uuid4

import docx
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from groq import Groq, GroqError
from PIL import Image
import fitz
import pytesseract
from pydantic import BaseModel
from pypdf import PdfReader
from qdrant_client import QdrantClient, models

from config import settings

app = FastAPI(title="RAG Backend", version="0.1.0")
frontend_base_url = settings.frontend_url

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        frontend_base_url
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


qdrant = QdrantClient(url=settings.qdrant_url)
groq_client = Groq(api_key=settings.groq_api_key)


@app.get("/health")
def health_check():
    return {
        "status": "ok",
        "groq_model": settings.groq_model,
        "qdrant_collection": settings.qdrant_collection,
        "qdrant_url": settings.qdrant_url,
    }

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
MIN_TEXT_LENGTH_PER_PAGE = 10
 
 
def ocr_image_bytes(image_bytes: bytes) -> str:
    """Runs OCR on raw image bytes and returns the recognized text."""
    image = Image.open(io.BytesIO(image_bytes))
    return pytesseract.image_to_string(image)
 
 
def ocr_pdf_page(page: fitz.Page) -> str:
    """Rasterizes a single PDF page to an image and runs OCR on it. Used as
    a fallback for pages that have no extractable text layer, e.g. a page
    that's really just a scanned image of a document."""
    pixmap = page.get_pixmap(dpi=300)
    image = Image.open(io.BytesIO(pixmap.tobytes("png")))
    return pytesseract.image_to_string(image)


def read_file_text(filename: str, content: bytes) -> str:
    """Extracts plain text from a PDF, DOCX, TXT, or image file's raw bytes.
 
    For PDFs, each page is checked individually: if pypdf finds a real text
    layer, that's used (fast, accurate). If a page comes back empty or
    near-empty, it's assumed to be a scanned image and gets OCR'd instead.
    This means a PDF with some normal pages and some scanned pages (e.g. a
    signed contract with one scanned signature page) is handled correctly
    without OCR'ing pages that didn't need it.
    """
    extension = Path(filename).suffix.lower()
 
    if extension == ".txt":
        return content.decode("utf-8", errors="replace")
 
    if extension in IMAGE_EXTENSIONS:
        return ocr_image_bytes(content)
 
    if extension == ".pdf":
        reader = PdfReader(io.BytesIO(content))
        ocr_doc = None 
        pages = []
 
        for page_number, page in enumerate(reader.pages):
            text = page.extract_text() or ""
 
            if len(text.strip()) < MIN_TEXT_LENGTH_PER_PAGE:
                if ocr_doc is None:
                    ocr_doc = fitz.open(stream=content, filetype="pdf")
                text = ocr_pdf_page(ocr_doc[page_number])
 
            pages.append(text)
 
        if ocr_doc is not None:
            ocr_doc.close()
 
        return "\n\n".join(pages)
 
    if extension == ".docx":
        document = docx.Document(io.BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        return "\n".join(paragraphs)
 
    raise HTTPException(status_code=400, detail=f"Unsupported file type: {filename}")


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 150) -> List[str]:
    """Splits text into overlapping chunks so a sentence cut off at one
    chunk boundary still appears in full in the next one."""
    text = text.strip()
    if not text:
        return []

    if len(text) <= chunk_size:
        return [text]

    chunks: List[str] = []
    start = 0
    step = chunk_size - overlap
    while start < len(text):
        chunk = text[start : start + chunk_size].strip()
        if chunk:
            chunks.append(chunk)
        start += step
    return chunks


def ensure_collection() -> None:
    """Creates the Qdrant collection on first use, sized to match whichever
    embedding model is configured, instead of a hardcoded vector size."""
    if not qdrant.collection_exists(settings.qdrant_collection):
        vector_size = qdrant.get_embedding_size(settings.embedding_model)
        qdrant.create_collection(
            collection_name=settings.qdrant_collection,
            vectors_config=models.VectorParams(
                size=vector_size, distance=models.Distance.COSINE
            ),
        )


@app.post("/ingest")
async def ingest(files: List[UploadFile] = File(...)):
    """Accepts one or more PDF/DOCX/TXT files, extracts and chunks their
    text, embeds each chunk locally via FastEmbed, and stores it in Qdrant.

    Note: Groq is not involved here. Groq has no embeddings API, so
    embeddings are computed locally and for free via FastEmbed (bundled
    with qdrant-client) instead. Groq comes in later, for generating answers
    in the /query endpoint.
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded")

    ensure_collection()

    results = []
    total_chunks = 0

    for upload in files:
        content = await upload.read()
        text = read_file_text(upload.filename, content)
        chunks = chunk_text(text)

        if not chunks:
            results.append(
                {
                    "filename": upload.filename,
                    "chunks_indexed": 0,
                    "note": "no extractable text found",
                }
            )
            continue

        points = [
            models.PointStruct(
                id=str(uuid4()),
                vector=models.Document(text=chunk, model=settings.embedding_model),
                payload={
                    "text": chunk,
                    "source": upload.filename,
                    "chunk_index": index,
                },
            )
            for index, chunk in enumerate(chunks)
        ]
        qdrant.upsert(collection_name=settings.qdrant_collection, points=points)

        results.append({"filename": upload.filename, "chunks_indexed": len(chunks)})
        total_chunks += len(chunks)

    return {
        "status": "success",
        "files": results,
        "total_chunks": total_chunks,
        "collection": settings.qdrant_collection,
    }


class QueryRequest(BaseModel):
    question: str
    top_k: int = 5


SYSTEM_PROMPT = (
    "You are a helpful assistant that answers questions using only the "
    "provided context. If the context doesn't contain enough information "
    "to answer, say so plainly instead of guessing or using outside "
    "knowledge. Refer to context sections by their [number] when you use "
    "them."
)


def search_chunks(question: str, top_k: int) -> List[dict]:
    """Embeds the question (via the same FastEmbed model used at ingest
    time) and returns the top_k most similar chunks from Qdrant, each with
    its source filename, position, and similarity score."""
    results = qdrant.query_points(
        collection_name=settings.qdrant_collection,
        query=models.Document(text=question, model=settings.embedding_model),
        limit=top_k,
    ).points

    return [
        {
            "text": point.payload["text"],
            "source": point.payload["source"],
            "chunk_index": point.payload["chunk_index"],
            "score": point.score,
        }
        for point in results
    ]


def build_context(chunks: List[dict]) -> str:
    """Formats retrieved chunks into one labeled block so the model (and the
    system prompt) can refer to each one by number."""
    sections = [
        f"[{i}] Source: {chunk['source']}\n{chunk['text']}"
        for i, chunk in enumerate(chunks, start=1)
    ]
    return "\n\n".join(sections)


def generate_answer(question: str, context: str) -> str:
    """Sends the retrieved context and the question to Groq and returns the
    generated answer."""
    try:
        response = groq_client.chat.completions.create(
            model=settings.groq_model,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {
                    "role": "user",
                    "content": f"Context:\n{context}\n\nQuestion: {question}",
                },
            ],
            temperature=0.2,
        )
    except GroqError as error:
        raise HTTPException(status_code=502, detail=f"Groq request failed: {error}")

    return response.choices[0].message.content


@app.post("/query")
def query(request: QueryRequest):
    """Searches Qdrant for chunks relevant to the question, then asks Groq
    to answer using only those chunks as context.

    This is a plain `def`, not `async def`, on purpose: both the Qdrant and
    Groq SDK calls here are blocking (synchronous) network calls. FastAPI
    runs plain `def` endpoints in a worker thread automatically, so a slow
    Groq response doesn't freeze the event loop for other requests. Marking
    this `async def` without using async clients would do the opposite of
    what you want.
    """
    if not settings.groq_api_key:
        raise HTTPException(status_code=500, detail="GROQ_API_KEY is not configured")

    if not qdrant.collection_exists(settings.qdrant_collection):
        raise HTTPException(
            status_code=400, detail="No documents have been ingested yet"
        )

    chunks = search_chunks(request.question, request.top_k)

    if not chunks:
        return {
            "answer": "I couldn't find anything relevant in the indexed documents.",
            "sources": [],
        }

    context = build_context(chunks)
    answer = generate_answer(request.question, context)

    return {
        "answer": answer,
        "sources": [
            {
                "source": chunk["source"],
                "chunk_index": chunk["chunk_index"],
                "score": round(chunk["score"], 4),
            }
            for chunk in chunks
        ],
    }